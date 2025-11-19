from restaurant import Dish, Menu, Order, Table, Database
from docx import Document
from openpyxl import Workbook
import os

class RestaurantApp:
    """Основной класс приложения ресторана"""
    
    def __init__(self):
        self.menu = Menu()
        self.tables = []
        self.database = Database()
        self._initialize_sample_data()
    
    def _initialize_sample_data(self):
        """Инициализация примеров данных"""
        # Создание примеров блюд
        dishes_data = [
            ("Пицца Маргарита", 450, 15, {"тесто": 50, "томаты": 30, "сыр": 40}),
            ("Паста Карбонара", 380, 12, {"паста": 40, "бекон": 35, "сливки": 25}),
            ("Салат Цезарь", 320, 8, {"салат": 20, "курица": 30, "сухарики": 10}),
            ("Стейк", 780, 20, {"мясо": 120, "специи": 15, "овощи": 25}),
            ("Тирамису", 280, 5, {"бисквит": 20, "маскарпоне": 25, "кофе": 10})
        ]
        
        for name, price, time, ingredients in dishes_data:
            dish = Dish(name, price, time, ingredients)
            self.menu.add_dish(dish)
        
        # Создание столиков
        for i in range(1, 6):
            seats = 4 if i % 2 == 0 else 2
            table = Table(i, seats)
            self.tables.append(table)
    
    def display_menu(self):
        """Отображение меню"""
        print("\n=== МЕНЮ РЕСТОРАНА ===")
        for i, dish in enumerate(self.menu.get_all_dishes(), 1):
            print(f"{i}. {dish}")
    
    def display_tables(self):
        """Отображение статуса столиков"""
        print("\n=== СТОЛИКИ ===")
        for table in self.tables:
            print(table)
    
    def create_order(self):
        """Создание нового заказа"""
        self.display_tables()
        
        try:
            table_id = int(input("\nВыберите номер столика: "))
            table = next((t for t in self.tables if t.table_id == table_id), None)
            
            if not table:
                print("Столик не найден!")
                return
            
            if table.is_occupied:
                print("Этот столик уже занят!")
                return
            
            table.occupy()
            order = table.get_current_order()
            
            while True:
                self.display_menu()
                choice = input("\nВыберите номер блюда (0 - завершить заказ): ")
                
                if choice == '0':
                    break
                
                try:
                    dish_index = int(choice) - 1
                    dishes = self.menu.get_all_dishes()
                    
                    if 0 <= dish_index < len(dishes):
                        dish = dishes[dish_index]
                        order.add_dish(dish)
                        print(f"Добавлено: {dish.name}")
                    else:
                        print("Неверный номер блюда!")
                
                except ValueError:
                    print("Введите корректный номер!")
            
            # Применение скидки
            discount_choice = input("Применить скидку? (y/n): ").lower()
            if discount_choice == 'y':
                try:
                    discount = float(input("Размер скидки (%): "))
                    order.apply_discount(discount)
                except ValueError:
                    print("Неверный формат скидки!")
            
            # Сохранение в базу данных
            self.database.save_order(order)
            
                       # Вывод итогов
            print(f"\n=== ЗАКАЗ ЗАВЕРШЕН ===")
            print(f"Столик: #{table_id}")
            
            if order.discount > 0:
                print(f"Общая сумма: {order.calculate_original_total():.2f} руб.")
                print(f"Скидка: {order.discount:.1f}% (-{order.get_discount_amount():.2f} руб.)")
                print(f"Итого к оплате: {order.calculate_total():.2f} руб.")
            else:
                print(f"Итого к оплате: {order.calculate_total():.2f} руб.")
                
            print(f"Время ожидания: {order.calculate_waiting_time()} мин.")
            print(f"Количество блюд: {len(order.get_dishes())}")
        
        except ValueError:
            print("Ошибка ввода данных!")
    
    def generate_doc_report(self):
        """Генерация отчета в формате DOCX"""
        doc = Document()
        doc.add_heading('Отчет по заказам ресторана', 0)
        
        orders = self.database.get_all_orders()
        
        if not orders:
            doc.add_paragraph('Нет данных о заказах')
        else:
            doc.add_paragraph(f'Всего заказов: {len(orders)}')
            
            for order in orders:
                doc.add_heading(f'Заказ #{order["order_id"]}', level=1)
                doc.add_paragraph(f'Дата: {order["order_date"]}')
                doc.add_paragraph(f'Сумма: {order["total_amount"]:.2f} руб.')
                doc.add_paragraph(f'Скидка: {order["discount"]}%')
                doc.add_paragraph(f'Время ожидания: {order["waiting_time"]} мин.')
                doc.add_paragraph(f'Блюда: {order["dishes"]}')
        
        filename = "reports/orders_report.docx"
        os.makedirs("reports", exist_ok=True)
        doc.save(filename)
        print(f"Отчет сохранен в {filename}")
    
    def generate_excel_report(self):
        """Генерация отчета в формате XLSX"""
        wb = Workbook()
        ws = wb.active
        ws.title = "Заказы"
        
        # Заголовки
        headers = ['ID заказа', 'Дата', 'Сумма', 'Скидка %', 'Время ожидания', 'Блюда']
        ws.append(headers)
        
        orders = self.database.get_all_orders()
        
        for order in orders:
            ws.append([
                order['order_id'],
                order['order_date'],
                order['total_amount'],
                order['discount'],
                order['waiting_time'],
                order['dishes']
            ])
        
        filename = "reports/orders_report.xlsx"
        os.makedirs("reports", exist_ok=True)
        wb.save(filename)
        print(f"Отчет сохранен в {filename}")
    
    def run(self):
        """Запуск основного цикла приложения"""
        while True:
            print("\n=== СИСТЕМА УПРАВЛЕНИЯ РЕСТОРАНОМ ===")
            print("1. Показать меню")
            print("2. Показать столики")
            print("3. Создать заказ")
            print("4. Сгенерировать отчет (DOCX)")
            print("5. Сгенерировать отчет (XLSX)")
            print("6. Выход")
            
            choice = input("Выберите действие: ")
            
            if choice == '1':
                self.display_menu()
            elif choice == '2':
                self.display_tables()
            elif choice == '3':
                self.create_order()
            elif choice == '4':
                self.generate_doc_report()
            elif choice == '5':
                self.generate_excel_report()
            elif choice == '6':
                print("До свидания!")
                break
            else:
                print("Неверный выбор!")

if __name__ == "__main__":
    app = RestaurantApp()
    app.run()